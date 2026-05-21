import io
import uuid
from typing import List

import pdfplumber
from docx import Document

from .models import CanonicalDoc, Block, Token, SourceSpan
from .utils import canonicalize

PARSER_VERSION = "docx-pdf-v1"


def _tokens_from_text(text: str, base_span: SourceSpan) -> List[Token]:
    tokens = []
    for idx, part in enumerate(text.split()):
        token_id = f"t-{uuid.uuid4().hex[:8]}-{idx}"
        tokens.append(Token(token_id=token_id, text=part, span=base_span))
    return tokens


def parse_docx(data: bytes) -> CanonicalDoc:
    doc = Document(io.BytesIO(data))
    blocks: List[Block] = []
    tokens: List[Token] = []

    for p_idx, para in enumerate(doc.paragraphs):
        text = canonicalize(para.text)
        if not text:
            continue
        span = SourceSpan(paragraph=p_idx)
        block_id = f"p-{p_idx}"
        blocks.append(Block(block_id=block_id, block_type="paragraph", text=text, span=span))
        for r_idx, run in enumerate(para.runs):
            r_text = canonicalize(run.text)
            if not r_text:
                continue
            r_span = SourceSpan(paragraph=p_idx, run=r_idx)
            tokens.extend(_tokens_from_text(r_text, r_span))

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = canonicalize(cell.text)
                if not text:
                    continue
                span = SourceSpan(table=t_idx, row=r_idx, col=c_idx)
                block_id = f"tbl-{t_idx}-{r_idx}-{c_idx}"
                blocks.append(Block(block_id=block_id, block_type="table_cell", text=text, span=span))
                tokens.extend(_tokens_from_text(text, span))

    token_map = {t.token_id: t.span for t in tokens}
    return CanonicalDoc(doc_id=f"doc-{uuid.uuid4().hex}", blocks=blocks, tokens=tokens, token_map=token_map)


def parse_pdf(data: bytes) -> CanonicalDoc:
    blocks: List[Block] = []
    tokens: List[Token] = []

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for p_idx, page in enumerate(pdf.pages):
            text = canonicalize(page.extract_text() or "")
            if text:
                span = SourceSpan(page=p_idx)
                block_id = f"pdf-{p_idx}"
                blocks.append(Block(block_id=block_id, block_type="paragraph", text=text, span=span))
                tokens.extend(_tokens_from_text(text, span))

            tables = page.extract_tables() or []
            for t_idx, table in enumerate(tables):
                for r_idx, row in enumerate(table):
                    for c_idx, cell in enumerate(row):
                        if not cell:
                            continue
                        ctext = canonicalize(str(cell))
                        span = SourceSpan(page=p_idx, table=t_idx, row=r_idx, col=c_idx)
                        block_id = f"pdf-tbl-{p_idx}-{t_idx}-{r_idx}-{c_idx}"
                        blocks.append(Block(block_id=block_id, block_type="table_cell", text=ctext, span=span))
                        tokens.extend(_tokens_from_text(ctext, span))

    token_map = {t.token_id: t.span for t in tokens}
    return CanonicalDoc(doc_id=f"doc-{uuid.uuid4().hex}", blocks=blocks, tokens=tokens, token_map=token_map)


def parse_upload(name: str, data: bytes) -> CanonicalDoc:
    name = (name or "").lower()
    if name.endswith(".docx"):
        return parse_docx(data)
    if name.endswith(".pdf"):
        return parse_pdf(data)
    # Fallback plain text
    raw = data.decode("utf-8", errors="ignore")
    blocks: List[Block] = []
    tokens: List[Token] = []
    for idx, line in enumerate(raw.splitlines()):
        line_text = canonicalize(line)
        if not line_text:
            continue
        span = SourceSpan(paragraph=idx)
        block_id = f"txt-{idx}"
        blocks.append(Block(block_id=block_id, block_type="paragraph", text=line_text, span=span))
        tokens.extend(_tokens_from_text(line_text, span))
    token_map = {t.token_id: t.span for t in tokens}
    return CanonicalDoc(doc_id=f"doc-{uuid.uuid4().hex}", blocks=blocks, tokens=tokens, token_map=token_map)
